#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将README.md文件转换为Word文档格式
"""

import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def markdown_to_docx(md_file, docx_file):
    """将Markdown文件转换为Word文档"""
    
    # 读取Markdown文件，确保UTF-8编码
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        # 如果UTF-8失败，尝试GBK编码
        with open(md_file, 'r', encoding='gbk') as f:
            content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体，确保中文显示正确
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    # 设置中文字体
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    font.size = Pt(10.5)
    
    # 逐行处理内容
    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()  # 保留内部空格，只去掉行尾空格
        
        if not line:  # 空行
            doc.add_paragraph()
            continue
            
        # 处理标题
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line[2:])
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        # 处理列表项
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:])
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        elif line.startswith('1. ') or re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            # 移除数字编号，让Word自动处理
            text = re.sub(r'^\d+\.\s', '', line)
            run = p.add_run(text)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        # 处理代码块
        elif line.startswith('```'):
            continue
            
        # 处理引用
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.italic = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
        # 处理普通段落
        else:
            # 处理粗体和斜体
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 移除**粗体**标记
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # 移除*斜体*标记
            
            # 检查是否是问题标题
            if text.startswith('Q: '):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            elif text.startswith('A: '):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.italic = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                if text.strip():  # 只有非空文本才添加段落
                    p = doc.add_paragraph(text)
                    for run in p.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 保存文档
    doc.save(docx_file)
    print(f"文档已成功转换并保存为: {docx_file}")

if __name__ == "__main__":
    import sys
    
    # 检查命令行参数
    if len(sys.argv) >= 2:
        md_file = sys.argv[1]
    else:
        md_file = "README.md"
    
    if len(sys.argv) >= 3:
        docx_file = sys.argv[2]
    else:
        # 默认输出文件名
        if md_file == "README.md":
            docx_file = "系统监控工具使用说明.docx"
        else:
            # 从输入文件名推导输出文件名
            base_name = md_file.rsplit('.', 1)[0]
            docx_file = f"{base_name}.docx"
    
    markdown_to_docx(md_file, docx_file)